import html
import json
import os
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str(Path("matplotlib_cache").resolve()))

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


TARGET_RESULT = Path("jobs/2026-06-16__19-22-34/result.json")
COMPARISON_RESULT = Path("jobs/2026-06-16__19-41-52/result.json")
FAILING_DETAILS = Path("jobs/2026-06-16__19-22-34/commitments-tracking__NAeqmhk/verifier/details.json")
OUTPUT_DOCX = Path("apollo_writeup.docx")
OUTPUT_HTML = Path("apollo_writeup.html")
FIGURE_1 = Path("figure_1_pass_rate.png")
FIGURE_2 = Path("figure_2_failure_counts.png")


# Load a JSON file from disk.
def load_json(path):
    with open(path, encoding="utf-8") as file:
        return json.load(file)


# Return the single eval block from a Harbor result.
def eval_block(result):
    return next(iter(result["stats"]["evals"].values()))


# Count deterministic passes from reward_stats.
def deterministic_summary(result):
    block = eval_block(result)
    buckets = block["reward_stats"]["deterministic_reward"]
    passed = len(buckets.get("1", []))
    total = block["n_trials"]
    return passed, total, passed / total


# Count failed trials per deterministic check.
def failure_counts(result):
    block = eval_block(result)
    ignored = {"reward", "deterministic_reward", "llm_judge_score", "llm_judge_available"}
    counts = {}
    for check, buckets in block["reward_stats"].items():
        if check not in ignored:
            counts[check] = len(buckets.get("0", []))
    return counts


# Gather all write-up metrics from the job files.
def compute_metrics():
    target = load_json(TARGET_RESULT)
    comparison = load_json(COMPARISON_RESULT)
    target_passed, target_total, target_rate = deterministic_summary(target)
    comparison_passed, comparison_total, comparison_rate = deterministic_summary(comparison)
    target_failures = failure_counts(target)
    nonzero_target_failures = {k: v for k, v in target_failures.items() if v}
    target_metrics = eval_block(target)["metrics"][0]
    comparison_metrics = eval_block(comparison)["metrics"][0]
    failing_details = load_json(FAILING_DETAILS)
    return {
        "target_passed": target_passed,
        "target_total": target_total,
        "target_rate": target_rate,
        "comparison_passed": comparison_passed,
        "comparison_total": comparison_total,
        "comparison_rate": comparison_rate,
        "target_failures": target_failures,
        "nonzero_target_failures": nonzero_target_failures,
        "target_metrics": target_metrics,
        "comparison_metrics": comparison_metrics,
        "failing_details": failing_details,
    }


# Create the pass-rate figure.
def make_pass_rate_figure(metrics):
    labels = ["Gemma 4 26B", "Gemma 4 31B"]
    rates = [metrics["target_rate"] * 100, metrics["comparison_rate"] * 100]
    counts = [f"{metrics['target_passed']}/{metrics['target_total']}", f"{metrics['comparison_passed']}/{metrics['comparison_total']}"]

    plt.figure(figsize=(6.5, 3.8))
    bars = plt.bar(labels, rates, color=["#4F81BD", "#70AD47"])
    plt.ylim(0, 110)
    plt.ylabel("Deterministic pass rate (%)")
    plt.title("Commitments-tracking pass rate by model")
    plt.grid(axis="y", alpha=0.25)
    for bar, rate, count in zip(bars, rates, counts):
        plt.text(bar.get_x() + bar.get_width() / 2, rate + 3, f"{count}\n{rate:.0f}%", ha="center", fontsize=10)
    plt.tight_layout()
    plt.savefig(FIGURE_1, dpi=180)
    plt.close()


# Create the target-model failure-count figure.
def make_failure_figure(metrics):
    failures = metrics["nonzero_target_failures"]
    ordered = sorted(failures.items(), key=lambda item: (-item[1], item[0]))
    labels = [name.replace("_", "\n") for name, _ in ordered]
    counts = [count for _, count in ordered]

    plt.figure(figsize=(6.5, 3.8))
    bars = plt.bar(range(len(labels)), counts, color="#C00000")
    plt.ylim(0, max(3, max(counts) + 1 if counts else 1))
    plt.ylabel("Failed trials out of 10")
    plt.title("Gemma 4 26B failed checks")
    plt.xticks(range(len(labels)), labels, fontsize=10)
    plt.grid(axis="y", alpha=0.25)
    for bar, count in zip(bars, counts):
        if count:
            plt.text(bar.get_x() + bar.get_width() / 2, count + 0.08, str(count), ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(FIGURE_2, dpi=180)
    plt.close()


# Apply the business brief style.
def set_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


# Add a paragraph to the document.
def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    return paragraph


# Add a compact bullet to the document.
def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


# Add a figure and caption to the document.
def add_figure(document, path, caption, width=6.2):
    document.add_picture(str(path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    caption_paragraph.paragraph_format.space_after = Pt(12)


# Add the title block.
def add_title(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Apollo Commitments-Tracking Eval")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Client update and loss analysis")
    subrun.italic = True
    subrun.font.size = Pt(11)


# Add the client update document.
def add_client_update(document, metrics):
    document.add_heading("Client Update", level=1)
    add_paragraph(
        document,
        "We built a Harbor knowledge-work task for evaluating whether an agent can keep a task board accurate from workplace conversation. "
        "The environment contains simulated Slack and Task Manager services exposed through terminal tools, a realistic instruction, a deterministic verifier, and an oracle solution. "
        "Everything runs in a containerized Harbor task, with the agent using the same service commands a human operator would use.",
    )
    add_paragraph(
        document,
        "The task asks the agent to identify outstanding commitments in Slack, reconcile them with the existing board, create missing tasks, assign the right owner, and avoid tasks for decoys such as FYIs, completed work, hypotheticals, and declined asks. "
        "The verifier writes a binary deterministic reward based on final state: the required commitments must be present with correct owners and actionable statuses, duplicates must be absent, and decoys must remain untracked.",
    )
    add_paragraph(
        document,
        f"The target model, Gemma 4 26B through terminus-2, passed {metrics['target_passed']} of {metrics['target_total']} trials for a deterministic pass rate of {metrics['target_rate']:.0%}. "
        f"The comparison model, Gemma 4 31B, passed {metrics['comparison_passed']} of {metrics['comparison_total']} trials for a deterministic pass rate of {metrics['comparison_rate']:.0%}. "
        "This makes the task challenging for the target model while remaining solvable for a stronger model.",
    )
    add_figure(document, FIGURE_1, "Figure 1. Deterministic pass rate by model.")
    add_paragraph(
        document,
        "The result is meaningful because the weaker model reliably handled most explicit commitments but intermittently dropped the billing API monitoring follow-up, a threaded and less direct obligation. "
        "The stronger model captured the full required end state across all recorded trials. That separation indicates the task discriminates on commitment inference and reconciliation rather than on basic command execution.",
    )
    add_paragraph(
        document,
        "Recommendation: use this eval as a compact regression test for operational commitment tracking. "
        "The next calibration step should preserve deterministic state checking while adding one or two more reconciliation cases only if future target models again approach saturation.",
    )


# Add the loss analysis document.
def add_loss_analysis(document, metrics):
    document.add_page_break()
    document.add_heading("Loss Analysis", level=1)
    add_paragraph(
        document,
        "The target model failed when it produced a task board that looked orderly but was incomplete. "
        "In one failing 26B trajectory, the agent created five of the six genuine missing commitments and then declared the board complete. "
        "The omitted item was the billing API monitoring follow-up assigned to Devon.",
    )
    add_paragraph(document, "The failing trajectory created these tasks:")
    for item in [
        "Update onboarding FAQ with new SSO steps -> alice",
        "Rotate staging API key -> priya",
        "Reconcile April invoice mismatch -> nina",
        "Draft Initech renewal email -> maya",
        "Complete Globex security questionnaire -> alice",
    ]:
        add_bullet(document, item)
    add_paragraph(
        document,
        "There was no create_task call for billing API monitoring. The deterministic verifier therefore marked billing_monitoring_ok = 0 and expected_task_count = 0, causing the all-or-nothing reward to fail the trial.",
    )
    add_figure(document, FIGURE_2, "Figure 2. Failed deterministic checks for Gemma 4 26B.")
    add_paragraph(
        document,
        f"Across the full target run, the only non-zero failed checks were billing_monitoring_ok ({metrics['nonzero_target_failures'].get('billing_monitoring_ok', 0)} failed trials) and expected_task_count ({metrics['nonzero_target_failures'].get('expected_task_count', 0)} failed trials). "
        "The stronger 31B model reached 10 of 10 deterministic passes, which shows that the environment is solvable and the target failure is not caused by an impossible verifier.",
    )
    add_paragraph(
        document,
        "The verification design deliberately treats deterministic state checking as authoritative. Success is objective here: either the task board contains the required commitments with the correct owners and no spurious work, or it does not. "
        "That makes deterministic reward more appropriate than a judge-only approach.",
    )
    add_paragraph(
        document,
        "The LLM judge remains useful as a comparison signal but is too permissive to drive reward. In the failing 26B trial, gpt-4o-mini scored the run 1.0 and claimed all commitments were tracked even though billing API monitoring was missing. "
        "That false pass is exactly why the deterministic verifier is the sole reward.",
    )
    add_paragraph(
        document,
        "The calibration arc also matters. The initial design saturated at 100% for both tested Gemma models. We hardened it with decoys plus implicit and threaded commitments, which made the task unsaturated for the target model and discriminating against the stronger comparison. "
        "During calibration, two unfair-verifier bugs were caught and fixed: an exact-match false negative on staging key rotation and a cancellation check that inverted model ranking. Those fixes illustrate the unfair-verifier risk called out in the brief and are part of making the final result trustworthy.",
    )
    add_paragraph(
        document,
        "Harbor framing: the submitted task includes instruction.md, a Dockerized environment with Slack and Task Manager CLIs, a deterministic verifier that writes reward.json, and an oracle solution. "
        "The recorded runs were executed with harbor run using terminus-2 and Gemma models through OpenRouter.",
    )


# Write a simple HTML copy with embedded image references.
def write_html(metrics):
    html_text = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Apollo Commitments-Tracking Eval</title>
  <style>
    body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; line-height: 1.5; color: #222; }}
    h1 {{ color: #1f4e79; }}
    h2 {{ color: #2e74b5; margin-top: 28px; }}
    img {{ max-width: 100%; border: 1px solid #ddd; }}
    .caption {{ font-style: italic; color: #555; text-align: center; }}
    li {{ margin-bottom: 4px; }}
  </style>
</head>
<body>
  <h1>Apollo Commitments-Tracking Eval</h1>
  <p><em>Client update and loss analysis</em></p>
  <h2>Client Update</h2>
  <p>We built a Harbor knowledge-work task for evaluating whether an agent can keep a task board accurate from workplace conversation. The environment contains simulated Slack and Task Manager services exposed through terminal tools, a realistic instruction, a deterministic verifier, and an oracle solution.</p>
  <p>The task asks the agent to identify outstanding commitments in Slack, reconcile them with the existing board, create missing tasks, assign the right owner, and avoid tasks for decoys such as FYIs, completed work, hypotheticals, and declined asks. The verifier writes a binary deterministic reward based on final state.</p>
  <p>The target model, Gemma 4 26B through terminus-2, passed {metrics['target_passed']} of {metrics['target_total']} trials for a deterministic pass rate of {metrics['target_rate']:.0%}. The comparison model, Gemma 4 31B, passed {metrics['comparison_passed']} of {metrics['comparison_total']} trials for a deterministic pass rate of {metrics['comparison_rate']:.0%}.</p>
  <img src="{html.escape(str(FIGURE_1))}" alt="Pass rate by model">
  <p class="caption">Figure 1. Deterministic pass rate by model.</p>
  <p>The weaker model reliably handled most explicit commitments but intermittently dropped the billing API monitoring follow-up. The stronger model captured the full required end state across all recorded trials.</p>
  <p>Recommendation: use this eval as a compact regression test for operational commitment tracking, preserving deterministic reward as the authoritative success signal.</p>
  <h2>Loss Analysis</h2>
  <p>In one failing 26B trajectory, the agent created five of six genuine missing commitments and omitted billing API monitoring. It created onboarding -> alice, staging key -> priya, invoice -> nina, Initech -> maya, and Globex -> alice.</p>
  <img src="{html.escape(str(FIGURE_2))}" alt="26B failed checks">
  <p class="caption">Figure 2. Failed deterministic checks for Gemma 4 26B.</p>
  <p>The only non-zero target failure counts were billing_monitoring_ok ({metrics['nonzero_target_failures'].get('billing_monitoring_ok', 0)}) and expected_task_count ({metrics['nonzero_target_failures'].get('expected_task_count', 0)}). The gpt-4o-mini judge falsely scored that failing trial 1.0, so deterministic state checking remains the sole reward.</p>
  <p>The stronger model captured all six genuine missing commitments in all ten trials. This contrast suggests that the failure is not generic tool use, but the harder act of noticing and preserving a threaded or implicit operational commitment.</p>
  <p>The calibration arc moved from a saturated 100%/100% design to an unsaturated, discriminating task after adding decoys plus implicit and threaded commitments. Two unfair-verifier bugs were caught and fixed during calibration: staging-key exact matching and a cancellation check that inverted model ranking.</p>
  <p>Harbor framing: the submitted task includes instruction.md, a Dockerized environment with Slack and Task Manager CLIs, a deterministic verifier that writes reward.json, and an oracle solution. The recorded runs used harbor run with terminus-2 and Gemma over OpenRouter.</p>
</body>
</html>
"""
    OUTPUT_HTML.write_text(html_text, encoding="utf-8")


# Generate figures, DOCX, and HTML.
def main():
    metrics = compute_metrics()
    print(f"Gemma 4 26B: {metrics['target_passed']}/{metrics['target_total']} deterministic passes = {metrics['target_rate']:.2f}")
    print(f"Gemma 4 31B: {metrics['comparison_passed']}/{metrics['comparison_total']} deterministic passes = {metrics['comparison_rate']:.2f}")
    print(f"26B non-zero failed checks: {metrics['nonzero_target_failures']}")

    make_pass_rate_figure(metrics)
    make_failure_figure(metrics)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    set_styles(document)
    add_title(document)
    add_client_update(document, metrics)
    add_loss_analysis(document, metrics)
    document.save(OUTPUT_DOCX)
    write_html(metrics)


# Run the write-up generator.
if __name__ == "__main__":
    main()
